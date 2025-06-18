import os
import requests
import json
from PIL import Image
from datetime import datetime
import asyncio
from notion_client import AsyncClient # Changed to AsyncClient
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re # Import re for regex
import argparse
from dotenv import load_dotenv
from notion_to_gdoc import upload_docx_to_gdoc
load_dotenv()


# Configuration (can be overridden by CLI arguments)
NOTION_API_TOKEN = os.getenv("NOTION_API_TOKEN")
DATABASE_ID = "165d5037135a807d9278d0d3c01e738a"
OUTPUT_FILENAME = "NotionContent.docx"
FILTER_HISTORY_FILE = "notion_filter_history.json"
DB_HISTORY_FILE = "notion_db_history.json"

# Initialize Notion client (will be initialized in main based on args)
notion_client_instance = None

def load_db_history():
    if os.path.exists(DB_HISTORY_FILE):
        with open(DB_HISTORY_FILE, 'r') as f:
            try:
                return json.load(f)
            except json.JSONDecodeError:
                return []
    return []

def save_db_history(db_ids):
    # Keep only the last 3 database IDs
    db_ids = db_ids[-3:]
    with open(DB_HISTORY_FILE, 'w') as f:
        json.dump(db_ids, f, indent=2)


def load_filter_history(file_path):
    if os.path.exists(file_path):
        with open(file_path, 'r') as f:
            try:
                return json.load(f)
            except json.JSONDecodeError:
                return []
    return []

def save_filter_history(filters, file_path):
    # Keep only the last 3 filters
    filters = filters[-3:]
    with open(file_path, 'w') as f:
        json.dump(filters, f, indent=2)


def format_filter_summary(filter_obj):
    parts = []
    if "and" in filter_obj:
        for sub_filter in filter_obj["and"]:
            parts.append(format_filter_summary(sub_filter))
        return "(" + " AND ".join(parts) + ")"
    elif "or" in filter_obj:
        for sub_filter in filter_obj["or"]:
            parts.append(format_filter_summary(sub_filter))
        return "(" + " OR ".join(parts) + ")"
    elif "property" in filter_obj:
        prop_name = filter_obj["property"]
        if "select" in filter_obj:
            return f"{prop_name} = {filter_obj['select']['equals']}"
        elif "status" in filter_obj:
            return f"{prop_name} = {filter_obj['status']['equals']}"
        elif "multi_select" in filter_obj:
            return f"{prop_name} CONTAINS {filter_obj['multi_select']['contains']}"
        elif "number" in filter_obj:
            num_filter = filter_obj["number"]
            if "equals" in num_filter:
                return f"{prop_name} = {num_filter['equals']}"
            elif "greater_than" in num_filter:
                return f"{prop_name} > {num_filter['greater_than']}"
            elif "less_than" in num_filter:
                return f"{prop_name} < {num_filter['less_than']}"
            elif "greater_than_or_equal_to" in num_filter:
                return f"{prop_name} >= {num_filter['greater_than_or_equal_to']}"
            elif "less_than_or_equal_to" in num_filter:
                return f"{prop_name} <= {num_filter['less_than_or_equal_to']}"
            elif "between" in num_filter:
                return f"{prop_name} BETWEEN {num_filter['greater_than_or_equal_to']} AND {num_filter['less_than_or_equal_to']}"
        elif "checkbox" in filter_obj:
            return f"{prop_name} = {filter_obj['checkbox']['equals']}"
    return "UNKNOWN_FILTER"


def add_rich_text_to_paragraph(paragraph, rich_texts):
    for rt in rich_texts:
        text_content = rt['plain_text']
        annotations = rt['annotations']
        
        run = paragraph.add_run(text_content)
        if annotations['bold']:
            run.bold = True
        if annotations['italic']:
            run.italic = True
        if annotations['strikethrough']:
            run.strike = True
        if annotations['underline']:
            run.underline = True
        if annotations['code']:
            run.font.name = 'Courier New'
            run.font.size = 10000

def create_checkbox(paragraph, checked):
    run = paragraph.add_run()
    if checked:
        run.add_text("☑ ")
    else:
        run.add_text("☐ ")
    run.font.name = 'Wingdings 2'
    run.font.size = 10000

def remove_excess_blank_lines(document):
    paragraphs_to_remove = []
    consecutive_blank_count = 0

    for paragraph in document.paragraphs:
        if paragraph.text.strip() == '':
            consecutive_blank_count += 1
            if consecutive_blank_count > 1:
                paragraphs_to_remove.append(paragraph)
        else:
            consecutive_blank_count = 0
    
    for p in paragraphs_to_remove:
        p._element.getparent().remove(p._element)

def extract_estimation_value(estimation_str):
    if not estimation_str or estimation_str == "N/A":
        return 0.0
    # Use regex to find numbers (integers or floats)
    match = re.search(r'(\d+(\.\d+)?)', estimation_str)
    if match:
        try:
            return float(match.group(1))
        except ValueError:
            return 0.0
    return 0.0

async def get_block_children(notion_client, block_id):
    response = await notion_client.blocks.children.list(block_id=block_id)
    return response['results']

async def process_blocks(notion_client, document, blocks, level=0):
    for block in blocks:
        block_type = block['type']
        
        if block_type == 'paragraph':
            paragraph = document.add_paragraph()
            add_rich_text_to_paragraph(paragraph, block['paragraph']['rich_text'])
        
        elif block_type.startswith('heading'):
            heading_level = int(block_type[-1])
            if heading_level == 1:
                paragraph = document.add_heading('', level=1)
            elif heading_level == 2:
                paragraph = document.add_heading('', level=2)
            elif heading_level == 3:
                paragraph = document.add_heading('', level=3)
            else:
                paragraph = document.add_paragraph(style='Normal')
            add_rich_text_to_paragraph(paragraph, block[block_type]['rich_text'])

        elif block_type == 'bulleted_list_item':
            paragraph = document.add_paragraph(style='List Bullet')
            paragraph.paragraph_format.left_indent = Inches(0.25 * level)
            add_rich_text_to_paragraph(paragraph, block['bulleted_list_item']['rich_text'])
            if block['has_children']:
                nested_blocks = await get_block_children(notion_client, block['id'])
                await process_blocks(notion_client, document, nested_blocks, level + 1)

        elif block_type == 'numbered_list_item':
            paragraph = document.add_paragraph(style='List Number')
            paragraph.paragraph_format.left_indent = Inches(0.25 * level)
            add_rich_text_to_paragraph(paragraph, block['numbered_list_item']['rich_text'])
            if block['has_children']:
                nested_blocks = await get_block_children(notion_client, block['id'])
                await process_blocks(notion_client, document, nested_blocks, level + 1)

        elif block_type == 'to_do':
            paragraph = document.add_paragraph()
            create_checkbox(paragraph, block['to_do']['checked'])
            add_rich_text_to_paragraph(paragraph, block['to_do']['rich_text'])

        elif block_type == 'image':
            image_url = block['image']['file']['url'] if 'file' in block['image'] else block['image']['external']['url']
            try:
                response = requests.get(image_url)
                response.raise_for_status()
                image_path = f"temp_image_{block['id']}.png"
                with open(image_path, 'wb') as f:
                    f.write(response.content)
                
                img = Image.open(image_path)
                original_width_px, original_height_px = img.size
                img.close()

                MAX_HEIGHT_INCHES = 2.75
                MAX_WIDTH_INCHES = 6.5
                DPI = 96

                original_width_inches = original_width_px / DPI
                original_height_inches = original_height_px / DPI

                target_width_inches = original_width_inches
                target_height_inches = original_height_inches

                if original_height_inches > MAX_HEIGHT_INCHES:
                    scale_factor = MAX_HEIGHT_INCHES / original_height_inches
                    target_height_inches = MAX_HEIGHT_INCHES
                    target_width_inches = original_width_inches * scale_factor

                if target_width_inches > MAX_WIDTH_INCHES:
                    scale_factor = MAX_WIDTH_INCHES / target_width_inches
                    target_width_inches = MAX_WIDTH_INCHES
                    target_height_inches = target_height_inches * scale_factor

                if target_width_inches != original_width_inches or target_height_inches != original_height_inches:
                    document.add_picture(image_path, width=Inches(target_width_inches), height=Inches(target_height_inches))
                else:
                    document.add_picture(image_path)
                
                os.remove(image_path)
            except requests.exceptions.RequestException as e:
                document.add_paragraph(f"Could not download image from {image_url}: {e}")
            except Exception as e:
                document.add_paragraph(f"Error processing image {image_url}: {e}")

        elif block_type == 'child_page':
            document.add_paragraph(f"--- Child Page: {block['child_page']['title']} ---")

        elif block_type == 'unsupported':
            document.add_paragraph(f"Unsupported block type: {block_type}")
        
        if block['has_children'] and block_type not in ['bulleted_list_item', 'numbered_list_item']:
            child_blocks = await get_block_children(notion_client, block['id'])
            await process_blocks(notion_client, document, child_blocks, level + 1)


def get_user_filters(filter_history, available_properties):
    property_filters = {}
    final_filter = {}

    if filter_history:
        print("\nPrevious filter configurations:")
        for i, hist_filter in enumerate(filter_history):
            print(f"{i+1}. {format_filter_summary(hist_filter)}")
        
        while True:
            choice = input(f"Enter number to use a previous filter, or 'n' for new filters: ").lower()
            if choice == 'n':
                break
            try:
                idx = int(choice) - 1
                if 0 <= idx < len(filter_history):
                    final_filter = filter_history[idx]
                    print(f"Using selected filter: {format_filter_summary(final_filter)}")
                    return final_filter
                else:
                    print("Invalid number. Please try again.")
            except ValueError:
                print("Invalid input. Please enter a number or 'n'.")
    
    while True:
        if property_filters:
            print("\nCurrent filters:")
            for prop_name, conditions in property_filters.items():
                if len(conditions) == 1:
                    print(f"- {prop_name}: {format_filter_summary(conditions[0])}")
                else:
                    print(f"- {prop_name} (OR conditions): {format_filter_summary({'or': conditions})}")
            print("-" * 40)

        add_filter_choice = input("Do you want to add another filter? (yes/no): ").lower()
        if add_filter_choice == 'no':
            all_combined_filters = []
            for prop_name, conditions in property_filters.items():
                if len(conditions) == 1:
                    all_combined_filters.append(conditions[0])
                else:
                    all_combined_filters.append({"or": conditions})
            
            if not all_combined_filters:
                print("No filters added. Proceeding without filters.")
                return {}
            
            print("\nSummary of all filters to be applied (AND-ed between properties, OR-ed within same property):")
            if all_combined_filters:
                if len(all_combined_filters) == 1:
                    print(f"- {format_filter_summary(all_combined_filters[0])}")
                else:
                    print(f"- {format_filter_summary({'and': all_combined_filters})}")
            else:
                print("- No filters applied.")
            
            confirm_filters = input("Confirm these filters? (yes/no): ").lower()
            if confirm_filters == 'yes':
                if all_combined_filters:
                    if len(all_combined_filters) == 1:
                        final_filter = all_combined_filters[0]
                    else:
                        final_filter = {"and": all_combined_filters}
                return final_filter
            else:
                print("Restarting filter selection...")
                property_filters = {}
                continue
        elif add_filter_choice != 'yes':
            print("Invalid choice. Please enter 'yes' or 'no'.")
            continue

        prop_name = input("Enter property name to filter on (e.g., 'Priority', 'STATUS', 'Estimation'): ")
        if prop_name not in available_properties:
            print(f"Error: Property '{prop_name}' not found. Please choose from available properties.")
            continue
        
        prop_type = available_properties[prop_name]
        filter_condition = {}

        if prop_type == 'select' or prop_type == 'status':
            values_str = input(f"Enter comma-separated values for '{prop_name}' (e.g., 'Mid, High', 'Refinement, Done'): ")
            values = [v.strip() for v in values_str.split(',')]
            for val in values:
                condition = {
                    "property": prop_name,
                    prop_type: {
                        "equals": val
                    }
                }
                if prop_name not in property_filters:
                    property_filters[prop_name] = []
                property_filters[prop_name].append(condition)

        elif prop_type == 'multi_select':
            values_str = input(f"Enter comma-separated values for '{prop_name}' (e.g., '10h, 20h'): ")
            values = [v.strip() for v in values_str.split(',')]
            for val in values:
                condition = {
                    "property": prop_name,
                    "multi_select": {
                        "contains": val
                    }
                }
                if prop_name not in property_filters:
                    property_filters[prop_name] = []
                property_filters[prop_name].append(condition)

        elif prop_type == 'number':
            op_type = input(f"Enter number filter type for '{prop_name}' (equals, greater_than, less_than, greater_than_or_equal_to, less_than_or_equal_to, between): ")
            if op_type == 'between':
                try:
                    start = float(input("Enter start value: "))
                    end = float(input("Enter end value: "))
                    filter_condition = {
                        "property": prop_name,
                        "number": {
                            "greater_than_or_equal_to": start,
                            "less_than_or_equal_to": end
                        }
                    }
                except ValueError:
                    print("Invalid number input. Please enter numeric values.")
                    continue
            else:
                try:
                    value = float(input(f"Enter value for '{prop_name}': "))
                    filter_condition = {
                        "property": prop_name,
                        "number": {
                            op_type: value
                        }
                    }
                except ValueError:
                    print("Invalid number input. Please enter a numeric value.")
                    continue
            
            if prop_name not in property_filters:
                property_filters[prop_name] = []
            property_filters[prop_name].append(filter_condition)

        elif prop_type == 'checkbox':
            value = input(f"Enter 'true' or 'false' for '{prop_name}': ").lower()
            if value == 'true':
                filter_condition = {"property": prop_name, "checkbox": {"equals": True}}
            elif value == 'false':
                filter_condition = {"property": prop_name, "checkbox": {"equals": False}}
            else:
                print("Invalid input for checkbox. Please enter 'true' or 'false'.")
                continue
            
            if prop_name not in property_filters:
                property_filters[prop_name] = []
            property_filters[prop_name].append(filter_condition)

        else:
            print(f"Filtering for property type '{prop_type}' is not yet supported.")
            continue
        
        print("Filter added.")

async def main():
    parser = argparse.ArgumentParser(description="Extract rich content from Notion database pages to a Word document.")
    parser.add_argument("--token", type=str, help="Notion API token.")
    parser.add_argument("--database_id", type=str, help="ID of the Notion database.")
    parser.add_argument("--output_file", type=str, default=OUTPUT_FILENAME,
                        help="Name of the output Word document file.")
    parser.add_argument("--document_name", type=str, 
                        help="Base name for the output Word document and Google Doc (e.g., 'MyNotionDoc'). Default is 'NotionContent'.")
    parser.add_argument("--filter_history_file", type=str, default=FILTER_HISTORY_FILE,
                        help="Path to the filter history file.")
    parser.add_argument("--db_history_file", type=str, default=DB_HISTORY_FILE,
                        help="Path to the database ID history file.")
    
    args = parser.parse_args()

    # Determine the base document name
    base_document_name = args.document_name
    if not base_document_name:
        base_document_name = input("Enter base name for the output document (e.g., 'MyNotionDoc', default: 'NotionContent'): ") or "NotionContent"

    timestamp = datetime.now().strftime("_%Y%m%d_%H%M")
    
    output_dir = "Output"
    os.makedirs(output_dir, exist_ok=True) # Ensure the Output directory exists

    args.output_file = os.path.join(output_dir, f"{base_document_name}{timestamp}.docx")

    notion_token = args.token
    if not notion_token:
        notion_token = os.environ.get("NOTION_API_TOKEN")
        if not notion_token and os.path.exists(".env"):
            try:
                with open(".env", 'r') as f:
                    for line in f:
                        if line.startswith("NOTION_API_TOKEN="):
                            notion_token = line.strip().split("=", 1)[1]
                            break
            except Exception as e:
                print(f"Warning: Could not read .env file: {e}")
        
        if not notion_token:
            notion_token = input("Enter Notion API Token (or set NOTION_API_TOKEN in .env or environment variables): ")
    
    if not notion_token:
        print("Notion API Token is required. Exiting.")
        return

    notion_client_instance = AsyncClient(auth=notion_token) # Changed to AsyncClient

    document = Document()
    
    current_time = datetime.now()
    formatted_time = current_time.strftime("%d-%m-%Y %H:%M")
    document.add_heading(f'Notion Database Content - Snapshot @ {formatted_time}', level=0)

    db_id = args.database_id
    db_history = load_db_history()
    if not db_id:
        if db_history:
            print("\nPrevious Database IDs:")
            for i, hist_db_id in enumerate(db_history):
                print(f"{i+1}. {hist_db_id}")
            while True:
                choice = input(f"Enter number to use a previous Database ID, 'n' for new, or leave blank for default ('{DATABASE_ID}'): ").lower()
                if choice == 'n':
                    db_id = input("Enter new Notion Database ID: ")
                    break
                elif not choice:
                    db_id = DATABASE_ID
                    break
                try:
                    idx = int(choice) - 1
                    if 0 <= idx < len(db_history):
                        db_id = db_history[idx]
                        print(f"Using selected Database ID: {db_id}")
                        break
                    else:
                        print("Invalid number. Please try again.")
                except ValueError:
                    print("Invalid input. Please enter a number, 'n', or leave blank.")
        else:
            db_id = input(f"Enter Notion Database ID (default: {DATABASE_ID}): ") or DATABASE_ID
    
    if not db_id:
        print("Database ID is required. Exiting.")
        return
    
    if db_id not in db_history:
        db_history.append(db_id)
        save_db_history(db_history)

    try:
        database_info = await notion_client_instance.databases.retrieve(database_id=db_id)
        available_properties = {}
        print("\nAvailable database properties for filtering:")
        for prop_name, prop_details in database_info['properties'].items():
            available_properties[prop_name] = prop_details['type']
            print(f"- {prop_name} (Type: {prop_details['type']})")
        print("-" * 40)
    except Exception as e:
        print(f"Error fetching database info: {e}")
        print("Cannot proceed without database property information. Please check database ID and token.")
        return

    filter_history = load_filter_history(args.filter_history_file)

    final_filter = get_user_filters(filter_history, available_properties)
    
    if final_filter and final_filter not in filter_history:
        filter_history.append(final_filter)
        save_filter_history(filter_history, args.filter_history_file)

    total_estimation_sum = 0.0 # Initialize total estimation sum

    try:
        query_params = {"database_id": db_id}
        if final_filter:
            query_params["filter"] = final_filter
        
        response = await notion_client_instance.databases.query(**query_params)
        pages = response['results']
        
        if not pages:
            document.add_paragraph("No pages found in the database matching your filters.")
            print("No pages found in the database matching your filters.")
            document.save(args.output_file)
            return

        for page in pages:
            page_id = page['id']
            page_title = "Untitled"
            if 'properties' in page:
                for prop_name, prop_value in page['properties'].items():
                    if prop_value['type'] == 'title' and prop_value['title']:
                        page_title = prop_value['title'][0]['plain_text']
                        break
            
            document.add_heading(f"Ticket: {page_title}", level=1)
            
            priority_name = "N/A"
            if 'Priority' in page['properties'] and page['properties']['Priority']['select']:
                priority_name = page['properties']['Priority']['select']['name']
            
            estimation_name = "N/A"
            if 'Estimation' in page['properties'] and page['properties']['Estimation']['multi_select']:
                estimation_name = ", ".join([item['name'] for item in page['properties']['Estimation']['multi_select']])
            
            # Extract and sum estimation value
            current_estimation_value = extract_estimation_value(estimation_name)
            total_estimation_sum += current_estimation_value

            subtitle_paragraph = document.add_paragraph()
            subtitle_run = subtitle_paragraph.add_run(f"Priority: {priority_name} | Estimation: {estimation_name}")
            subtitle_run.font.color.rgb = RGBColor(0x00, 0x00, 0x80)
            
            print(f"Processing ticket: {page_title} (ID: {page_id})")

            page_blocks = await get_block_children(notion_client_instance, page_id)
            await process_blocks(notion_client_instance, document, page_blocks)
            
            divider_paragraph = document.add_paragraph()
            divider_paragraph.add_run("--- END OF TICKET ---").bold = True
            divider_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        remove_excess_blank_lines(document)

        document.save(args.output_file)
        print(f"Successfully extracted Notion content to {args.output_file}")
        print(f"Total estimated hours for processed tickets: {total_estimation_sum:.2f}h") # Print total sum

        gdoc_name = f"{base_document_name}{timestamp}"
        print(f"Attempting to upload {args.output_file} to Google Docs as {gdoc_name}...")
        upload_docx_to_gdoc(args.output_file, gdoc_name)

    except Exception as e:
        print(f"An error occurred: {e}")
        document.add_paragraph(f"An error occurred during extraction: {e}")
        document.save(args.output_file)

if __name__ == "__main__":
    asyncio.run(main())
